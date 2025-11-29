import math
import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

from PIL import Image, ImageDraw, ImageFont

try:
    from docx import Document
    from docx.shared import Mm
except ImportError:
    Document = None
    Mm = None


# -----------------------------
#  Geometry / Rendering Helpers
# -----------------------------

def generate_box_image(board_width_mm, board_height_mm, finger_width_mm, dpi, start_with_finger):
    mm_to_px = dpi / 25.4
    margin_mm = 5
    margin_px = int(round(margin_mm * mm_to_px))

    board_width_px = int(round(board_width_mm * mm_to_px))
    board_height_px = int(round(board_height_mm * mm_to_px))

    img_width = board_width_px + 2 * margin_px
    img_height = board_height_px + 2 * margin_px + 80  # title + ruler

    img = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(img)

    # Fonts
    try:
        font_title = ImageFont.truetype("arial.ttf", 20)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except IOError:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()

    title = f"Box Joint Template - Board: {board_width_mm} mm, Finger: {finger_width_mm} mm"
    info = f"Print at {dpi} DPI, 'Actual size / 100%' with no scaling."
    draw.text((margin_px, margin_px), title, fill="black", font=font_title)
    draw.text((margin_px, margin_px + 24), info, fill="black", font=font_small)

    origin_y = margin_px + 50
    left_x = margin_px
    top_y = origin_y
    right_x = left_x + board_width_px
    bottom_y = top_y + board_height_px

    # Board outline
    draw.rectangle([left_x, top_y, right_x, bottom_y], outline="black", width=1)

    # Fingers
    x = left_x
    end_x = right_x
    draw_finger = start_with_finger
    finger_width_px = int(round(finger_width_mm * mm_to_px))

    while x < end_x:
        seg_width = min(finger_width_px, end_x - x)
        if draw_finger and seg_width > 0:
            draw.rectangle([x, top_y, x + seg_width, bottom_y], fill="black", outline=None)
        x += seg_width
        draw_finger = not draw_finger

    # Ruler
    ruler_top_y = bottom_y + 10
    ruler_length_mm = min(board_width_mm, 200)
    ruler_length_px = int(round(ruler_length_mm * mm_to_px))
    draw.line([left_x, ruler_top_y, left_x + ruler_length_px, ruler_top_y], fill="black", width=1)

    tick_step_mm = 10
    mm_val = 0
    while mm_val <= ruler_length_mm + 0.01:
        px = left_x + int(round(mm_val * mm_to_px))
        draw.line([px, ruler_top_y, px, ruler_top_y + 6], fill="black", width=1)
        label = f"{int(mm_val)} mm"
        draw.text((px - 10, ruler_top_y + 8), label, fill="black", font=font_small)
        mm_val += tick_step_mm

    return img


def parse_ratio(text: str) -> float:
    """
    Supports:
      "6"  -> 6.0 (1:6)
      "1:6" -> 6.0
    """
    text = text.strip()
    if ":" in text:
        parts = text.split(":")
        left = float(parts[0])
        right = float(parts[1])
        if left == 0:
            raise ValueError("Ratio left side cannot be zero.")
        return right / left
    else:
        return float(text)


def generate_dovetail_image(board_width_mm, board_height_mm, num_tails, ratio_val, dpi):
    mm_to_px = dpi / 25.4
    margin_mm = 5
    margin_px = int(round(margin_mm * mm_to_px))

    board_width_px = int(round(board_width_mm * mm_to_px))
    board_height_px = int(round(board_height_mm * mm_to_px))

    img_width = board_width_px + 2 * margin_px
    img_height = board_height_px + 2 * margin_px + 100

    img = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype("arial.ttf", 20)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except IOError:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()

    title = f"Dovetail Template - Board: {board_width_mm} mm, Tails: {num_tails}, Ratio: 1:{ratio_val:g}"
    info = f"Print at {dpi} DPI, 'Actual size / 100%' with no scaling."

    margin = margin_px
    draw.text((margin, margin), title, fill="black", font=font_title)
    draw.text((margin, margin + 24), info, fill="black", font=font_small)

    origin_y = margin + 50
    left_x = margin
    top_y = origin_y
    right_x = left_x + board_width_px
    bottom_y = top_y + board_height_px

    # Board outline
    draw.rectangle([left_x, top_y, right_x, bottom_y], outline="black", width=1)

    # Tails & pins layout
    pin_fraction = 0.5  # pin width relative to tail width
    tail_units = 1.0
    pin_units = pin_fraction
    total_units = num_tails * tail_units + (num_tails + 1) * pin_units
    unit_px = board_width_px / total_units

    tail_top_width_px = tail_units * unit_px
    pin_width_px = pin_units * unit_px
    offset_px = board_height_px / ratio_val

    tail_start_x = left_x + pin_width_px

    for i in range(num_tails):
        tail_left_top_x = tail_start_x + i * (tail_top_width_px + pin_width_px)
        tail_right_top_x = tail_left_top_x + tail_top_width_px

        tail_left_bottom_x = tail_left_top_x + offset_px
        tail_right_bottom_x = tail_right_top_x - offset_px

        if tail_right_bottom_x <= tail_left_bottom_x:
            mid = (tail_left_bottom_x + tail_right_bottom_x) / 2.0
            tail_left_bottom_x = mid
            tail_right_bottom_x = mid

        polygon = [
            (tail_left_top_x, top_y),
            (tail_right_top_x, top_y),
            (tail_right_bottom_x, bottom_y),
            (tail_left_bottom_x, bottom_y),
        ]
        draw.polygon(polygon, fill="black", outline=None)

    # Ruler
    ruler_top_y = bottom_y + 10
    ruler_length_mm = min(board_width_mm, 200)
    ruler_length_px = int(round(ruler_length_mm * mm_to_px))
    draw.line([left_x, ruler_top_y, left_x + ruler_length_px, ruler_top_y], fill="black", width=1)

    tick_step_mm = 10
    mm_val = 0
    while mm_val <= ruler_length_mm + 0.01:
        px = left_x + int(round(mm_val * mm_to_px))
        draw.line([px, ruler_top_y, px, ruler_top_y + 6], fill="black", width=1)
        label = f"{int(mm_val)} mm"
        draw.text((px - 10, ruler_top_y + 8), label, fill="black", font=font_small)
        mm_val += tick_step_mm

    return img


def save_box_svg(path, board_width_mm, board_height_mm, finger_width_mm, start_with_finger):
    """
    SVG in mm units, suitable for LightBurn import.
    """
    width = board_width_mm
    height = board_height_mm

    lines = []
    lines.append(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}mm" height="{height}mm" viewBox="0 0 {width} {height}">'
    )
    lines.append('<g stroke="black" fill="none" stroke-width="0.1">')

    # Outline
    lines.append(f'<rect x="0" y="0" width="{width}" height="{height}" />')

    # Fingers along width
    x = 0.0
    end_x = width
    draw_finger = start_with_finger

    while x < end_x - 1e-6:
        seg = min(finger_width_mm, end_x - x)
        if draw_finger and seg > 0:
            lines.append(f'<rect x="{x}" y="0" width="{seg}" height="{height}" />')
        x += seg
        draw_finger = not draw_finger

    lines.append("</g></svg>")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def save_dovetail_svg(path, board_width_mm, board_height_mm, num_tails, ratio_val):
    """
    SVG in mm units, suitable for LightBurn import.
    """
    width = board_width_mm
    height = board_height_mm

    pin_fraction = 0.5
    tail_units = 1.0
    pin_units = pin_fraction
    total_units = num_tails * tail_units + (num_tails + 1) * pin_units
    unit_mm = width / total_units

    tail_top_width_mm = tail_units * unit_mm
    pin_width_mm = pin_units * unit_mm
    offset_mm = height / ratio_val
    tail_start_x = pin_width_mm

    lines = []
    lines.append(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}mm" height="{height}mm" viewBox="0 0 {width} {height}">'
    )
    lines.append('<g stroke="black" fill="none" stroke-width="0.1">')

    # Outline
    lines.append(f'<rect x="0" y="0" width="{width}" height="{height}" />')

    for i in range(num_tails):
        left_top = tail_start_x + i * (tail_top_width_mm + pin_width_mm)
        right_top = left_top + tail_top_width_mm

        left_bottom = left_top + offset_mm
        right_bottom = right_top - offset_mm

        if right_bottom <= left_bottom:
            mid = (left_bottom + right_bottom) / 2.0
            left_bottom = mid
            right_bottom = mid

        points = [
            (left_top, 0),
            (right_top, 0),
            (right_bottom, height),
            (left_bottom, height),
        ]
        points_str = " ".join(f"{x},{y}" for x, y in points)
        lines.append(f'<polygon points="{points_str}" />')

    lines.append("</g></svg>")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# -----------------------------
#  GUI & Wiring
# -----------------------------

def update_field_states():
    jt = joint_type_var.get()
    if jt == "box":
        entry_finger_width.config(state="normal")
        chk_start_with_finger.config(state="normal")
        entry_tails.config(state="disabled")
        entry_ratio.config(state="disabled")
    else:
        entry_finger_width.config(state="disabled")
        chk_start_with_finger.config(state="disabled")
        entry_tails.config(state="normal")
        entry_ratio.config(state="normal")


def generate_templates():
    try:
        board_width_mm = float(entry_board_width.get())
        board_height_mm = float(entry_height.get())
        dpi = int(entry_dpi.get())

        if board_width_mm <= 0 or board_height_mm <= 0 or dpi <= 0:
            messagebox.showerror("Invalid input", "Board dimensions and DPI must be > 0.")
            return

        outputs_selected = any([
            var_out_png.get(),
            var_out_svg.get(),
            var_out_pdf.get(),
            var_out_docx.get()
        ])
        if not outputs_selected:
            messagebox.showerror("No output format", "Select at least one output format.")
            return

        joint_type = joint_type_var.get()

        # Joint-specific params
        if joint_type == "box":
            finger_width_mm = float(entry_finger_width.get())
            if finger_width_mm <= 0 or finger_width_mm > board_width_mm:
                messagebox.showerror("Invalid input", "Finger width must be > 0 and <= board width.")
                return
            start_with_finger = bool(var_start_with_finger.get())
            img = generate_box_image(board_width_mm, board_height_mm, finger_width_mm, dpi, start_with_finger)
        else:
            num_tails = int(entry_tails.get())
            if num_tails < 1:
                messagebox.showerror("Invalid input", "Number of tails must be at least 1.")
                return
            ratio_val = parse_ratio(entry_ratio.get())
            if ratio_val <= 0:
                messagebox.showerror("Invalid input", "Dovetail ratio must be > 0.")
                return
            img = generate_dovetail_image(board_width_mm, board_height_mm, num_tails, ratio_val, dpi)

        # Ask for base file name
        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG Image", "*.png"), ("All files", "*.*")],
            title="Save template (base name)"
        )
        if not file_path:
            return

        base, _ = os.path.splitext(file_path)
        messages = []

        # Always create PNG as the master image (even if user doesn't want PNG "output")
        png_path = base + ".png"
        img.save(png_path, "PNG")
        if var_out_png.get():
            messages.append(f"PNG: {png_path}")

        # PDF
        if var_out_pdf.get():
            pdf_path = base + ".pdf"
            img.save(pdf_path, "PDF", resolution=dpi)
            messages.append(f"PDF: {pdf_path}")

        # SVG (LightBurn)
        if var_out_svg.get():
            svg_path = base + ".svg"
            if joint_type == "box":
                save_box_svg(svg_path, board_width_mm, board_height_mm, finger_width_mm, start_with_finger)
            else:
                save_dovetail_svg(svg_path, board_width_mm, board_height_mm, num_tails, ratio_val)
            messages.append(f"SVG (LightBurn): {svg_path}")

        # Word / DOCX
        if var_out_docx.get():
            if Document is None or Mm is None:
                raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
            docx_path = base + ".docx"
            doc = Document()
            doc.add_heading("Joint Template", level=1)
            jt_label = "Box joint" if joint_type == "box" else "Dovetail joint"
            doc.add_paragraph(f"Type: {jt_label}")
            doc.add_paragraph(f"Board width: {board_width_mm} mm")
            doc.add_paragraph(f"Template height/depth: {board_height_mm} mm")
            if joint_type == "box":
                doc.add_paragraph(f"Finger width: {finger_width_mm} mm")
                doc.add_paragraph(f"Start with finger: {start_with_finger}")
            else:
                doc.add_paragraph(f"Tails: {num_tails}")
                doc.add_paragraph(f"Dovetail ratio: 1:{ratio_val:g}")
            # Insert picture roughly at full board width + margin
            doc.add_picture(png_path, width=Mm(board_width_mm + 10))
            doc.save(docx_path)
            messages.append(f"Word (DOCX): {docx_path}")

        messagebox.showinfo("Templates generated", "Generated:\n" + "\n".join(messages))

    except ValueError as ve:
        messagebox.showerror("Invalid input", str(ve))
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{e}")


# -----------------------------
#  Build GUI
# -----------------------------

root = tk.Tk()
root.title("Joint Template Generator")

mainframe = ttk.Frame(root, padding="10 10 10 10")
mainframe.grid(row=0, column=0, sticky="nsew")
root.columnconfigure(0, weight=1)
root.rowconfigure(0, weight=1)

# Joint type
joint_type_var = tk.StringVar(value="box")
ttk.Label(mainframe, text="Joint type:").grid(row=0, column=0, sticky="w")
ttk.Radiobutton(mainframe, text="Box joint", variable=joint_type_var, value="box",
                command=update_field_states).grid(row=0, column=1, sticky="w")
ttk.Radiobutton(mainframe, text="Dovetail", variable=joint_type_var, value="dovetail",
                command=update_field_states).grid(row=0, column=2, sticky="w")

# Common fields
ttk.Label(mainframe, text="Board width (mm):").grid(row=1, column=0, sticky="w", pady=2)
entry_board_width = ttk.Entry(mainframe, width=10)
entry_board_width.grid(row=1, column=1, sticky="w", pady=2)
entry_board_width.insert(0, "100")

ttk.Label(mainframe, text="Template height/depth (mm):").grid(row=2, column=0, sticky="w", pady=2)
entry_height = ttk.Entry(mainframe, width=10)
entry_height.grid(row=2, column=1, sticky="w", pady=2)
entry_height.insert(0, "40")

ttk.Label(mainframe, text="Output DPI:").grid(row=3, column=0, sticky="w", pady=2)
entry_dpi = ttk.Entry(mainframe, width=10)
entry_dpi.grid(row=3, column=1, sticky="w", pady=2)
entry_dpi.insert(0, "300")

# Box joint fields
ttk.Label(mainframe, text="Finger width (mm):").grid(row=4, column=0, sticky="w", pady=2)
entry_finger_width = ttk.Entry(mainframe, width=10)
entry_finger_width.grid(row=4, column=1, sticky="w", pady=2)
entry_finger_width.insert(0, "10")

var_start_with_finger = tk.IntVar(value=1)
chk_start_with_finger = ttk.Checkbutton(
    mainframe,
    text="Start with finger (else gap)",
    variable=var_start_with_finger
)
chk_start_with_finger.grid(row=4, column=2, sticky="w", pady=2)

# Dovetail fields
ttk.Label(mainframe, text="Number of tails:").grid(row=5, column=0, sticky="w", pady=2)
entry_tails = ttk.Entry(mainframe, width=10)
entry_tails.grid(row=5, column=1, sticky="w", pady=2)
entry_tails.insert(0, "3")

ttk.Label(mainframe, text="Dovetail ratio (1:x):").grid(row=6, column=0, sticky="w", pady=2)
entry_ratio = ttk.Entry(mainframe, width=10)
entry_ratio.grid(row=6, column=1, sticky="w", pady=2)
entry_ratio.insert(0, "6")  # 1:6 default

# Output format checkboxes
ttk.Label(mainframe, text="Output formats:").grid(row=7, column=0, sticky="w", pady=(8, 2))

var_out_png = tk.IntVar(value=1)
var_out_svg = tk.IntVar(value=1)
var_out_pdf = tk.IntVar(value=0)
var_out_docx = tk.IntVar(value=0)

ttk.Checkbutton(mainframe, text="PNG (printable)", variable=var_out_png).grid(row=7, column=1, sticky="w")
ttk.Checkbutton(mainframe, text="SVG (LightBurn)", variable=var_out_svg).grid(row=7, column=2, sticky="w")
ttk.Checkbutton(mainframe, text="PDF", variable=var_out_pdf).grid(row=8, column=1, sticky="w")
ttk.Checkbutton(mainframe, text="Word (.docx)", variable=var_out_docx).grid(row=8, column=2, sticky="w")

# Buttons
btn_generate = ttk.Button(mainframe, text="Generate", command=generate_templates)
btn_generate.grid(row=9, column=0, sticky="w", pady=10)

btn_close = ttk.Button(mainframe, text="Close", command=root.destroy)
btn_close.grid(row=9, column=2, sticky="e", pady=10)

# Initial field state
update_field_states()

root.mainloop()
